from flask import Flask, request, jsonify
from core_ai import generate_summary
import os
from PyPDF2 import PdfReader
from docx import Document
from werkzeug.utils import secure_filename
from flask_cors import CORS



app = Flask(__name__)
CORS(app)


@app.route('/summarize', methods=['POST'])
def summarize():
    data = request.json
    text = data.get('text', '')
    summary = generate_summary(text)
    return jsonify({'summary': summary})

@app.route('/')
def home():
    return "Welcome to the Smart Document Assistant API! Use the /summarize endpoint for text summarization."

@app.route('/favicon.ico')
def favicon():
    return '', 204  # Return an empty response with a 204 (No Content) status

@app.route('/generate-template', methods=['POST'])
def generate_template():
    """
    API endpoint to generate a document template based on user-provided placeholders.
    """
    data = request.json
    template_type = data.get('template_type', 'default')
    placeholders = data.get('placeholders', {})

    # Define sample templates
    templates = {
        'email': "Dear {name},\n\nThis is a reminder about {event} scheduled on {date}. Please let us know if you need any assistance.\n\nBest regards,\nYour Team",
        'report': "Project Name: {project_name}\n\nStatus: {status}\n\nDeadline: {deadline}\n\nSummary: {summary}",
        'default': "Hello {name}, this is your default template."
    }

    # Get the selected template
    template = templates.get(template_type, templates['default'])

    try:
        # Fill in placeholders
        filled_template = template.format(**placeholders)
        return jsonify({'template': filled_template})
    except KeyError as e:
        return jsonify({'error': f"Missing placeholder: {e}"}), 400

from transformers import pipeline

# Initialize the question-answering pipeline
qa_pipeline = pipeline("question-answering", model="distilbert-base-cased-distilled-squad")

@app.route('/query', methods=['POST'])
def query_document():
    """
    API endpoint to process a query against a document and return the most relevant answer.
    """
    # Check if request is JSON or form-data
    if request.is_json:
        data = request.json
        document = data.get('document', '')
        query = data.get('query', '')
    else:  # Handle form-data
        document = request.form.get('document', '')
        query = request.form.get('query', '')

    # Log the inputs
    app.logger.info(f"Received document: {document[:200]}...")  # Log first 200 chars of document
    app.logger.info(f"Received query: {query}")

    if not document or not query:
        return jsonify({"error": "Both 'document' and 'query' are required."}), 400

    try:
        # Use the QA pipeline to find the answer
        result = qa_pipeline(question=query, context=document)
        app.logger.info(f"QA result: {result}")
        return jsonify({"answer": result['answer'], "confidence": result['score']})
    except Exception as e:
        app.logger.error(f"Error in QA pipeline: {str(e)}")
        return jsonify({"error": str(e)}), 500

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


# Initialize the summarization pipeline
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

@app.route('/upload', methods=['POST'])
def upload_document():
    """
    API endpoint to upload, parse, and summarize a document (PDF or Word).
    """
    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request."}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({"error": "No file selected."}), 400

    # Save the uploaded file
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    try:
        # Parse text from the file
        text = ''
        if file.filename.endswith('.pdf'):
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text()
        elif file.filename.endswith('.docx'):
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + '\n'
        else:
            return jsonify({"error": "Unsupported file format. Only PDF and DOCX are allowed."}), 400

        # Summarize the extracted text
        if len(text.strip()) == 0:
            return jsonify({"error": "The uploaded document is empty or unreadable."}), 400

        # Handle large text by splitting it into chunks
        max_chunk_size = 1024
        text_chunks = [text[i:i + max_chunk_size] for i in range(0, len(text), max_chunk_size)]
        summarized_text = ""
        for chunk in text_chunks:
            summary = summarizer(chunk, max_length=130, min_length=30, do_sample=False)
            summarized_text += summary[0]['summary_text'] + " "

        return jsonify({"extracted_text": summarized_text.strip(), "summary": summarized_text.strip()}) #orginally extracted_text value is text.

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/query-upload', methods=['POST'])
def query_uploaded_document():
    if 'file' not in request.files or 'query' not in request.form:
        return jsonify({"error": "Both 'file' and 'query' are required."}), 400

    file = request.files['file']
    query = request.form['query']

    if file.filename == '':
        return jsonify({"error": "No file selected."}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    try:
        text = ''
        if file.filename.endswith('.pdf'):
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text()
        elif file.filename.endswith('.docx'):
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + '\n'
        else:
            return jsonify({"error": "Unsupported file format. Only PDF and DOCX are allowed."}), 400
        
        result = qa_pipeline(question=query, context=text)
        return jsonify({"answer": result['answer'], "confidence": result['score']})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(debug=True)
